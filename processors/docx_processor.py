"""Document processor for .docx files using python-docx."""

import logging
from pathlib import Path

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from .base import BaseDocumentProcessor, ProcessorResult

logger = logging.getLogger(__name__)

# Heading style name prefixes → markdown heading levels
_HEADING_MAP = {
    "Heading 1": "# ",
    "Heading 2": "## ",
    "Heading 3": "### ",
    "Heading 4": "#### ",
    "Heading 5": "##### ",
    "Heading 6": "###### ",
    "Title": "# ",
    "Subtitle": "## ",
}


def _table_to_markdown(table) -> str:
    """Convert a docx table to a markdown table."""
    rows = []
    for row in table.rows:
        cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
    if len(rows) >= 1:
        # Insert separator after header row
        col_count = len(table.rows[0].cells)
        separator = "| " + " | ".join(["---"] * col_count) + " |"
        rows.insert(1, separator)
    return "\n".join(rows)


def _is_list_paragraph(paragraph) -> tuple[bool, bool]:
    """Check if a paragraph is a list item.

    Returns (is_list, is_ordered).
    """
    style_name = (paragraph.style.name or "").lower()
    if "list bullet" in style_name:
        return True, False
    if "list number" in style_name:
        return True, True
    # Check for numPr element in paragraph properties (covers custom list styles)
    pPr = paragraph._element.find(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr"
    )
    if pPr is not None:
        numPr = pPr.find(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr"
        )
        if numPr is not None:
            return True, "number" in style_name
    return False, False


class DocxProcessor(BaseDocumentProcessor):
    """Extracts text from .docx files as structured markdown."""

    @property
    def supported_media_types(self) -> set[str]:
        return {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"}

    @property
    def supported_extensions(self) -> set[str]:
        return {"docx"}

    def process(self, file_path: Path, output_dir: Path) -> ProcessorResult:
        try:
            doc = Document(str(file_path))
        except (PackageNotFoundError, Exception) as e:
            return ProcessorResult(error=f"Failed to open .docx: {e}")

        lines: list[str] = []
        ordered_counter = 0

        # Interleave paragraphs and tables in document order
        body = doc.element.body
        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        # Build lookup maps
        para_elements = list(doc.paragraphs)
        table_elements = list(doc.tables)
        para_idx = 0
        table_idx = 0

        for child in body:
            tag = child.tag
            if tag == f"{ns}p" and para_idx < len(para_elements):
                para = para_elements[para_idx]
                para_idx += 1

                text = para.text.strip()
                if not text:
                    # Preserve blank lines for readability
                    if lines and lines[-1] != "":
                        lines.append("")
                    continue

                style_name = para.style.name or ""

                # Headings
                heading_prefix = None
                for key, prefix in _HEADING_MAP.items():
                    if style_name.startswith(key):
                        heading_prefix = prefix
                        break

                if heading_prefix:
                    if lines and lines[-1] != "":
                        lines.append("")
                    lines.append(f"{heading_prefix}{text}")
                    lines.append("")
                    ordered_counter = 0
                    continue

                # Lists
                is_list, is_ordered = _is_list_paragraph(para)
                if is_list:
                    if is_ordered:
                        ordered_counter += 1
                        lines.append(f"{ordered_counter}. {text}")
                    else:
                        lines.append(f"- {text}")
                        ordered_counter = 0
                    continue

                # Regular paragraph
                ordered_counter = 0
                lines.append(text)

            elif tag == f"{ns}tbl" and table_idx < len(table_elements):
                table = table_elements[table_idx]
                table_idx += 1
                if lines and lines[-1] != "":
                    lines.append("")
                lines.append(_table_to_markdown(table))
                lines.append("")
                ordered_counter = 0

        extracted_text = "\n".join(lines).strip() + "\n"

        # Write to .extracted.md alongside the original
        out_path = output_dir / f"{file_path.stem}.extracted.md"
        out_path.write_text(extracted_text, encoding="utf-8")

        return ProcessorResult(
            extracted_text=extracted_text,
            extracted_path=out_path,
            metadata={"paragraphs": para_idx, "tables": table_idx},
        )
