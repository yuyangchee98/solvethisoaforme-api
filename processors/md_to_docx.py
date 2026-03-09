"""Convert Markdown text to a DOCX document in memory."""

import re
from io import BytesIO

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _strip_backslash_escapes(text: str) -> str:
    r"""Remove backslash escapes like `1\.` → `1.` for markdown punctuation."""
    return re.sub(r"\\([\\`*_{}[\]()#+\-.!|>~])", r"\1", text)


def _apply_inline_formatting(paragraph, text: str):
    """Parse bold/italic markers and add formatted runs to a paragraph."""
    text = _strip_backslash_escapes(text)
    # Pattern matches **bold**, <u>underline</u>, ~~strikethrough~~, *italic*, or `code` segments
    pattern = re.compile(r"(\*\*(.+?)\*\*|<u>(.+?)</u>|~~(.+?)~~|\*(.+?)\*|`(.+?)`)")
    last_end = 0
    for m in pattern.finditer(text):
        # Add any plain text before this match
        if m.start() > last_end:
            run = paragraph.add_run(text[last_end : m.start()])
            run.font.size = Pt(11)
            run.font.name = "Calibri"
        if m.group(2):  # bold
            run = paragraph.add_run(m.group(2))
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = "Calibri"
        elif m.group(3):  # underline
            run = paragraph.add_run(m.group(3))
            run.underline = True
            run.font.size = Pt(11)
            run.font.name = "Calibri"
        elif m.group(4):  # strikethrough
            run = paragraph.add_run(m.group(4))
            run.font.strike = True
            run.font.size = Pt(11)
            run.font.name = "Calibri"
        elif m.group(5):  # italic
            run = paragraph.add_run(m.group(5))
            run.italic = True
            run.font.size = Pt(11)
            run.font.name = "Calibri"
        elif m.group(6):  # inline code
            run = paragraph.add_run(m.group(6))
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        last_end = m.end()
    # Trailing plain text
    if last_end < len(text):
        run = paragraph.add_run(text[last_end:])
        run.font.size = Pt(11)
        run.font.name = "Calibri"


def _strip_emsp_indent(text: str) -> tuple[str, int]:
    """Count and strip leading &emsp; entities, returning (stripped_text, indent_level)."""
    level = 0
    while text.startswith("&emsp;"):
        text = text[6:]  # len("&emsp;") == 6
        level += 1
    return text, level


def markdown_to_docx(markdown_text: str) -> BytesIO:
    """Convert markdown text to a DOCX file returned as an in-memory BytesIO.

    Supports headings (#-###), **bold**, <u>underline</u>, ~~strikethrough~~, *italic*, `inline code`,
    bullet lists (- item), numbered lists (1. item), fenced code blocks,
    &emsp; indentation, and plain paragraphs.
    """
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    lines = markdown_text.split("\n")
    i = 0
    in_code_block = False
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]

        # Fenced code block toggle
        if line.strip().startswith("```"):
            if in_code_block:
                # End code block — flush accumulated lines
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        stripped = line.strip()

        # Blank line → skip
        if not stripped:
            i += 1
            continue

        # Strip &emsp; indentation before matching patterns
        stripped, indent_level = _strip_emsp_indent(stripped)

        # Horizontal rule
        if re.match(r"^([*_-])\1{2,}\s*$", stripped):
            p = doc.add_paragraph()
            # Thin gray bottom border to simulate an HR
            pPr = p._element.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "999999")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Table (look-ahead: collect consecutive | lines)
        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            # Separate header, separator, and data rows
            if len(table_lines) >= 2 and re.match(r"^\|[\s:_-]+\|", table_lines[1]):
                header_cells = [c.strip() for c in table_lines[0].strip("|").split("|")]
                data_rows = []
                for tl in table_lines[2:]:
                    data_rows.append([c.strip() for c in tl.strip("|").split("|")])
                num_cols = len(header_cells)
                table = doc.add_table(rows=1 + len(data_rows), cols=num_cols, style="Table Grid")
                # Header row
                for ci, cell_text in enumerate(header_cells):
                    cell = table.rows[0].cells[ci]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    _apply_inline_formatting(p, cell_text)
                    for run in p.runs:
                        run.bold = True
                        run.font.name = "Calibri"
                        run.font.size = Pt(10)
                    # Light gray shading
                    shading = OxmlElement("w:shd")
                    shading.set(qn("w:fill"), "D9D9D9")
                    shading.set(qn("w:val"), "clear")
                    cell._element.get_or_add_tcPr().append(shading)
                # Data rows
                for ri, row_cells in enumerate(data_rows):
                    for ci in range(min(len(row_cells), num_cols)):
                        cell = table.rows[ri + 1].cells[ci]
                        cell.text = ""
                        p = cell.paragraphs[0]
                        _apply_inline_formatting(p, row_cells[ci])
                        for run in p.runs:
                            run.font.name = "Calibri"
                            run.font.size = Pt(10)
            else:
                # Not a real table, treat lines as paragraphs
                for tl in table_lines:
                    p = doc.add_paragraph()
                    _apply_inline_formatting(p, tl)
            continue

        # Blockquote (look-ahead: collect consecutive > lines)
        if stripped.startswith(">"):
            bq_parts = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                bq_line = re.sub(r"^>\s?", "", lines[i].strip())
                bq_parts.append(bq_line)
                i += 1
            bq_text = " ".join(bq_parts)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            _apply_inline_formatting(p, bq_text)
            for run in p.runs:
                run.italic = True
            continue

        # Headings — custom styled paragraphs (not Word heading styles)
        heading_match = re.match(r"^(#{1,6})\s+(.*)", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            p = doc.add_paragraph()
            _apply_inline_formatting(p, text)
            for run in p.runs:
                run.font.color.rgb = None  # black
                if level == 1:
                    run.font.size = Pt(14)
                    run.bold = True
                elif level == 2:
                    run.font.size = Pt(12)
                    run.bold = True
                elif level == 3:
                    run.font.size = Pt(12)
                    run.bold = True
                    run.underline = True
                else:  # 4-6
                    run.font.size = Pt(12)
                    run.underline = True
            i += 1
            continue

        # Bullet list item
        bullet_match = re.match(r"^[-*]\s+(.*)", stripped)
        if bullet_match:
            text = bullet_match.group(1)
            p = doc.add_paragraph(style="List Bullet")
            _apply_inline_formatting(p, text)
            i += 1
            continue

        # Numbered list item
        num_match = re.match(r"^\d+\.\s+(.*)", stripped)
        if num_match:
            text = num_match.group(1)
            p = doc.add_paragraph(style="List Number")
            _apply_inline_formatting(p, text)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        if indent_level > 0:
            p.paragraph_format.left_indent = Inches(0.5 * indent_level)
        _apply_inline_formatting(p, stripped)
        i += 1

    # Flush any unclosed code block
    if code_lines:
        p = doc.add_paragraph()
        run = p.add_run("\n".join(code_lines))
        run.font.name = "Courier New"
        run.font.size = Pt(9)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
